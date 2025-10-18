from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from news2docx.infra.logging import unified_print


def _sanitize_title_for_display(title: str) -> str:
    """Remove publisher suffix and trailing punctuation from titles.
    Example: "... hardware. | The Verge" -> "... hardware"
    - Drop anything after ' | '
    - Strip trailing punctuation like . ! ? 銆?锛?锛?"""
    if not title:
        return ""
    t = str(title).strip()
    if " | " in t:
        t = t.split(" | ", 1)[0].strip()
    t = t.rstrip(".?!銆傦紒锛?")
    return t


@dataclass
class FontConfig:
    name: str
    size_pt: float


@dataclass
class DocumentConfig:
    # First-line indent in inches (default 0.2 inch)
    first_line_indent_inch: float = 0.2
    font_zh: FontConfig = field(default_factory=lambda: FontConfig(name="宋体", size_pt=10.5))
    font_en: FontConfig = field(default_factory=lambda: FontConfig(name="Cambria", size_pt=10.5))
    title_size_multiplier: float = 1.0
    bilingual: bool = True
    order: str = "en-zh"  # English first


def _safe_title(s: str) -> str:
    s = (s or "").strip()
    return s if s else "Untitled"


def _split_paragraphs(text: str) -> List[str]:
    if not text:
        return []
    t = text.strip()
    if "%%" in t:
        parts = [p.strip() for p in t.split("%%") if p.strip()]
        return parts
    # fallback by blank lines
    parts = [p.strip() for p in re.split(r"\n\s*\n", t) if p.strip()]
    if parts:
        return parts
    # final fallback: split sentences for very short text
    sents = [p.strip() for p in re.split(r"(?<=[.!?])\s+", t) if p.strip()]
    return sents


def _strip_markdown(text: str, drop_headings: bool = True) -> str:
    """Remove common Markdown markers and optional heading lines.
    - Drops fenced code blocks and inline code markers
    - Replaces links [text](url) -> text; drops images ![alt](url)
    - Removes emphasis markers *, _, **, __, ~~ while keeping inner text
    - Optionally removes heading lines that start with '#'
    - Removes leading list/blockquote markers
    """
    if not text:
        return ""
    t = str(text)
    # Fenced code blocks
    t = re.sub(r"```[\s\S]*?```", "", t)
    # Inline code backticks
    t = re.sub(r"`([^`]*)`", r"\1", t)
    # Images: remove
    t = re.sub(r"!\[[^\]]*\]\([^\)]*\)", "", t)
    # Links: keep text
    t = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", t)
    # Emphasis
    t = re.sub(r"(\*\*|__)(.*?)\1", r"\2", t)
    t = re.sub(r"(\*|_)(.*?)\1", r"\2", t)
    t = re.sub(r"~~(.*?)~~", r"\1", t)
    # Blockquote and list markers at line start
    t = re.sub(r"^\s*>\s?", "", t, flags=re.MULTILINE)
    t = re.sub(r"^\s*(?:[-*+]\s+|\d+\.\s+)", "", t, flags=re.MULTILINE)
    if drop_headings:
        lines = []
        for line in t.splitlines():
            if re.match(r"^\s*#{1,6}\s+", line):
                continue
            lines.append(line)
        t = "\n".join(lines)
    return t


class DocumentWriter:
    def __init__(self, cfg: DocumentConfig | None = None) -> None:
        self.cfg = cfg or DocumentConfig()

    def _apply_font(self, run, zh: bool) -> None:
        if zh:
            run.font.name = self.cfg.font_zh.name
            run.font.size = Pt(self.cfg.font_zh.size_pt)
            # Ensure East Asia font set to 瀹嬩綋
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "瀹嬩綋")
            except Exception:
                pass
        else:
            run.font.name = self.cfg.font_en.name
            run.font.size = Pt(self.cfg.font_en.size_pt)

    def _write_title(self, doc: Document, title_text: str, zh: bool) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            _strip_markdown(
                _sanitize_title_for_display(_safe_title(title_text)), drop_headings=True
            )
        )
        # Apply font with title size multiplier
        if zh:
            r.font.name = self.cfg.font_zh.name
            r.font.size = Pt(
                self.cfg.font_zh.size_pt * float(self.cfg.title_size_multiplier or 1.0)
            )
            try:
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
        else:
            r.font.name = self.cfg.font_en.name
            r.font.size = Pt(
                self.cfg.font_en.size_pt * float(self.cfg.title_size_multiplier or 1.0)
            )
        r.font.bold = bool(getattr(self.cfg, "title_bold", True))

    def _write_block(self, doc: Document, paras: List[str], zh: bool) -> None:
        prev_text: str | None = None
        for para in paras:
            # Normalize to avoid hard line breaks inside a paragraph
            text = _strip_markdown(para, drop_headings=True)
            text = re.sub(r"[\r\n]+", " ", text).strip()
            if not text:
                continue
            # Deduplicate consecutive identical paragraphs
            if prev_text is not None and text == prev_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(self.cfg.first_line_indent_inch)
            r = p.add_run(text)
            self._apply_font(r, zh=zh)
            prev_text = text

    def write_from_processed(self, processed: Dict[str, Any], output_path: str) -> str:
        """Write a single DOCX containing all articles (legacy combined export)."""
        articles = processed.get("articles", []) if isinstance(processed, dict) else []
        if not articles:
            raise ValueError("Empty processed result; cannot export DOCX")

        doc = Document()

        for idx, art in enumerate(articles, start=1):
            zh_title = art.get("translated_title") or art.get("original_title") or ""
            en_title = art.get("original_title") or ""
            en_content = art.get("adjusted_content") or art.get("original_content") or ""
            zh_content = art.get("translated_content") or ""

            en_paras = _split_paragraphs(en_content)
            zh_paras = _split_paragraphs(zh_content)

            if self.cfg.bilingual and self.cfg.order == "en-zh":
                self._write_title(doc, en_title, zh=False)
                self._write_block(doc, en_paras, zh=False)
                self._write_title(doc, zh_title, zh=True)
                self._write_block(doc, zh_paras, zh=True)
            elif self.cfg.bilingual:
                self._write_title(doc, zh_title, zh=True)
                self._write_block(doc, zh_paras, zh=True)
                self._write_title(doc, en_title, zh=False)
                self._write_block(doc, en_paras, zh=False)
            else:
                self._write_title(doc, zh_title, zh=True)
                self._write_block(doc, zh_paras, zh=True)

            if idx < len(articles):
                doc.add_page_break()

        unified_print(f"Exported DOCX: {output_path}", "export", "docx")
        doc.save(output_path)
        return output_path

    def write_per_article(
        self, processed: Dict[str, Any], output_dir: str | None = None
    ) -> List[str]:
        """Write one DOCX per article with required formatting and naming by Chinese title."""
        articles = processed.get("articles", []) if isinstance(processed, dict) else []
        if not articles:
            raise ValueError("Empty processed result; cannot export DOCX")

        base = Path(output_dir) if output_dir else Path(".")
        base.mkdir(parents=True, exist_ok=True)
        out_paths: List[str] = []
        used: set[str] = set()

        def _filename_from_title(name: str) -> str:
            s = (_sanitize_title_for_display(_safe_title(name)) or "Untitled").strip()
            for ch in '\\/:*?"<>|':
                s = s.replace(ch, " ")
            s = " ".join(s.split())
            s = s[:80]
            stem = s or "Untitled"
            candidate = stem
            i = 1
            while candidate.lower() in used:
                i += 1
                candidate = f"{stem}-{i}"
            used.add(candidate.lower())
            return candidate + ".docx"

        for art in articles:
            zh_title = art.get("translated_title") or art.get("original_title") or ""
            en_title = art.get("original_title") or ""
            en_content = art.get("adjusted_content") or art.get("original_content") or ""
            zh_content = art.get("translated_content") or ""

            doc = Document()
            # English first
            self._write_title(doc, en_title, zh=False)
            self._write_block(doc, _split_paragraphs(en_content), zh=False)
            self._write_title(doc, zh_title, zh=True)
            self._write_block(doc, _split_paragraphs(zh_content), zh=True)

            fn = _filename_from_title(zh_title)
            path = base / fn
            doc.save(str(path))
            unified_print(f"Exported DOCX: {path}", "export", "docx")
            out_paths.append(str(path))

        return out_paths
