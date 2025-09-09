#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
新闻爬虫与文档生成主程序
直接执行完整的新闻处理流程：爬取 -> AI处理 -> 文档生成
合并了main.py、news_pipeline.py和docx_writer.py的所有功能

集成Scraper.py模块：Scraper.py提供纯爬虫功能，News2Docx_CLI.py提供AI处理功能
"""

# 标准库导入
import sys
import os
import re
import json
import time
import uuid
import random
import logging
import argparse
import glob
import itertools
from dataclasses import dataclass, asdict, field
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Tuple, Any, Union
from urllib.parse import urlparse
from pathlib import Path

# 第三方库导入
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 从Scraper.py导入爬虫相关功能
from Scraper import (
    ScrapeConfig, ScrapeResults,
    HttpClient, ContentExtractor, NewsAPIService, URLStore,
    NewsProcessingError, ScrapingError, APIError,
    now_stamp, ensure_directory,
    build_arg_parser, save_scraped_data_to_json
)
from Scraper import NewsScraper as ScraperNewsScraper

# 从ai_processor.py导入AI处理相关功能
from ai_processor import (
    Article, ProcessingError, DocumentError,
    estimate_max_tokens, build_source_block, build_system_and_user_prompt,
    backoff_sleep, call_siliconflow_tagged_output,  # 修正函数名
    generate_summary_single_article, repair_summary_if_too_short,
    process_single_article, split_articles_blocks, parse_paragraphs,
    validate_and_adjust_word_count, build_final_json_struct,
    process_articles_concurrent, process_articles_with_two_steps,
    process_articles_to_words, count_english_words, safe_filename,
    handle_error, safe_execute, now_stamp,
    DEFAULT_MODEL_ID, DEFAULT_SILICONFLOW_API_KEY, DEFAULT_SILICONFLOW_URL,  # 修正常量名
    DEFAULT_BATCH_SIZE, DEFAULT_MAX_TOKENS_HARD_CAP, DEFAULT_USE_TWO_STEP,
    DEFAULT_CONCURRENCY, CHINESE_VARIANT_HINT, SEPARATOR_LINE, MAIN_SEPARATOR,
    TAG_RE, PARA_TAG_RE,
    # 新增的两步处理函数
    process_articles_two_steps_concurrent, TARGET_WORD_MIN, TARGET_WORD_MAX
)

# 统一日志系统导入
from unified_logger import (
    get_unified_logger, log_task_start, log_task_end,
    log_processing_step, log_error, log_performance,
    unified_print, log_processing_result, log_file_operation,
    log_batch_processing
)


# -------------------------------
# 数据格式转换函数
# -------------------------------

def convert_ai_processor_output_to_docx_format(processed_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    将ai_processor的输出转换为News2Docx_CLI期望的格式

    Args:
        processed_data: ai_processor的输出数据

    Returns:
        Dict[str, Any]: 转换为News2Docx_CLI期望格式的数据
    """
    if not processed_data or "articles" not in processed_data:
        return processed_data

    converted_articles = []

    for article in processed_data["articles"]:
        try:
            # 转换单篇文章
            converted_article = convert_single_article_for_docx(article)
            if converted_article:
                converted_articles.append(converted_article)
        except Exception as e:
            unified_print(f"转换文章 {article.get('id', 'unknown')} 时出错: {e}", "news2docx_cli", "format_conversion", "warning")
            continue

    # 返回转换后的数据结构
    return {
        "articles": converted_articles,
        "metadata": processed_data.get("metadata", {}),
        "total": len(converted_articles),
        "success": len(converted_articles),
        "failed": len(processed_data.get("articles", [])) - len(converted_articles)
    }


def convert_single_article_for_docx(article: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    将单篇文章从ai_processor格式转换为News2Docx_CLI格式

    Args:
        article: 单篇文章数据（ai_processor格式）

    Returns:
        Optional[Dict[str, Any]]: 转换后的文章数据，转换失败返回None
    """
    try:
        article_id = article.get("id", "unknown")

        # 获取英文标题
        eng_title = article.get("original_title", "").strip()

        # 获取英文内容
        eng_content_raw = article.get("adjusted_content", article.get("original_content", ""))

        # 获取翻译后的中文标题和内容
        translated_title = article.get("translated_title", "").strip()
        translated_content = article.get("translated_content", "")

        # 如果没有翻译内容，尝试使用原始内容
        if not translated_content:
            chi_title = translated_title if translated_title else eng_title  # 使用翻译标题或英文标题作为中文标题
            chi_content_raw = eng_content_raw
        else:
            # 使用翻译后的标题，如果没有则从翻译内容中提取
            if translated_title:
                chi_title = translated_title
                chi_content_raw = translated_content
            else:
                # 尝试从翻译内容中提取中文标题（通常是第一行）
                lines = translated_content.split('\n', 1)
                chi_title = lines[0].strip() if lines else eng_title
                # 中文内容是除标题外的其余部分
                chi_content_raw = lines[1] if len(lines) > 1 else translated_content

        # 按段落分割英文内容
        eng_content_dict = split_content_into_paragraphs(eng_content_raw, "english")

        # 按段落分割中文内容
        chi_content_dict = split_content_into_paragraphs(chi_content_raw, "chinese")

        # 构建转换后的文章数据
        converted_article = {
            "id": article_id,
            "EngTitle": eng_title,
            "ChiTitle": chi_title,
            "EngContent": eng_content_dict,
            "ChiContent": chi_content_dict,
            "url": article.get("url", ""),
            "processing_timestamp": article.get("processing_timestamp", ""),
            "target_language": article.get("target_language", "Chinese"),
            "success": article.get("success", True)
        }

        return converted_article

    except Exception as e:
        unified_print(f"转换文章时出错: {e}", "news2docx_cli", "format_conversion", "error")
        return None


def split_content_into_paragraphs(content: str, language: str = "english") -> Dict[str, str]:
    """
    将文本内容按段落分割成字典格式

    Args:
        content: 原始文本内容
        language: 语言类型 ("english" 或 "chinese")

    Returns:
        Dict[str, str]: 段落字典 {"p1": "content1", "p2": "content2", ...}
    """
    if not content:
        return {}

    # 按换行符分割段落
    paragraphs = [p.strip() for p in content.split('\n') if p.strip()]

    # 如果没有换行符，按句子分割（对于英文）
    if len(paragraphs) == 1 and language == "english":
        # 英文按句号分割
        sentences = re.split(r'(?<=[.!?])\s+', content)
        paragraphs = [s.strip() for s in sentences if s.strip()]

    # 如果还是只有一个段落，检查是否有%%分隔符
    if len(paragraphs) == 1:
        if '%%' in content:
            paragraphs = [p.strip() for p in content.split('%%') if p.strip()]

    # 构建段落字典
    paragraph_dict = {}
    for i, para in enumerate(paragraphs, 1):
        if para:  # 只添加非空段落
            paragraph_dict[f"p{i}"] = para

    return paragraph_dict


def convert_scraper_output_to_docx_format(scraper_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    将Scraper的输出转换为News2Docx_CLI期望的格式

    Args:
        scraper_data: Scraper的输出数据

    Returns:
        Dict[str, Any]: 转换为News2Docx_CLI期望格式的数据
    """
    if not scraper_data or "articles" not in scraper_data:
        return scraper_data

    converted_articles = []

    for article in scraper_data["articles"]:
        try:
            # 转换单篇文章
            converted_article = convert_scraper_article_for_docx(article)
            if converted_article:
                converted_articles.append(converted_article)
        except Exception as e:
            unified_print(f"转换Scraper文章 {article.get('index', 'unknown')} 时出错: {e}", "news2docx_cli", "format_conversion", "warning")
            continue

    # 返回转换后的数据结构
    return {
        "articles": converted_articles,
        "metadata": {
            "total": scraper_data.get("total", 0),
            "success": scraper_data.get("success", 0),
            "failed": scraper_data.get("failed", 0),
            "processing_method": "scraper_only",
            "target_language": "none"
        },
        "total": scraper_data.get("total", 0),
        "success": scraper_data.get("success", 0),
        "failed": scraper_data.get("failed", 0)
    }


def convert_scraper_article_for_docx(article: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    将Scraper单篇文章转换为News2Docx_CLI格式

    Args:
        article: Scraper文章数据

    Returns:
        Optional[Dict[str, Any]]: 转换后的文章数据
    """
    try:
        article_id = str(article.get("index", "unknown"))

        # 获取标题和内容
        eng_title = article.get("title", "").strip()
        eng_content_raw = article.get("content", "")

        # 对于Scraper输出，没有中文内容，使用英文作为默认
        chi_title = eng_title
        chi_content_raw = eng_content_raw

        # 按段落分割内容
        eng_content_dict = split_content_into_paragraphs(eng_content_raw, "english")
        chi_content_dict = split_content_into_paragraphs(chi_content_raw, "chinese")

        # 构建转换后的文章数据
        converted_article = {
            "id": article_id,
            "EngTitle": eng_title,
            "ChiTitle": chi_title,
            "EngContent": eng_content_dict,
            "ChiContent": chi_content_dict,
            "url": article.get("url", ""),
            "scraped_at": article.get("scraped_at", ""),
            "content_length": article.get("content_length", 0),
            "word_count": article.get("word_count", 0),
            "success": True
        }

        return converted_article

    except Exception as e:
        unified_print(f"转换Scraper文章时出错: {e}", "news2docx_cli", "format_conversion", "error")
        return None


# -------------------------------
# 文档配置（News2Docx专用）
# -------------------------------

# 文档配置
DEFAULT_FIRST_LINE_INDENT = 0.3
DEFAULT_FONT_ZH_NAME = "宋体"
DEFAULT_FONT_EN_NAME = "Cambria"
# 字体大小配置（五号字体）
DEFAULT_FONT_SIZE_PT = 10.5  # 五号字体 = 10.5pt
DEFAULT_TITLE_SIZE_MULTIPLIER = 1.0  # 标题和正文一样大，都是10.5pt

# -------------------------------
# 数据类定义（News2Docx专用）
# -------------------------------

def create_output_filename(prefix: str = "", suffix: str = "", extension: str = "") -> str:
    """创建输出文件名"""
    parts = []
    if prefix:
        parts.append(prefix)
    parts.append(now_stamp())
    if suffix:
        parts.append(suffix)

    filename = "_".join(parts)
    if extension and not extension.startswith('.'):
        extension = f".{extension}"

    return f"{filename}{extension}"


# -------------------------------
# 工具类定义
# -------------------------------

# -------------------------------
# 爬虫主体
# -------------------------------




# -------------------------------
# CLI
# -------------------------------




def run_news_pipeline(cfg: ScrapeConfig) -> ScrapeResults:
    """
    运行完整的新闻处理管道：爬取 -> AI处理 -> 保存结果

    Args:
        cfg: 配置对象

    Returns:
        ScrapeResults: 爬取结果
    """
    # 记录任务开始
    log_task_start("news2docx_cli", "pipeline_processing", {
        "max_urls": cfg.max_urls,
        "concurrency": cfg.concurrency,
        "target_word_range": f"{TARGET_WORD_MIN}-{TARGET_WORD_MAX}"
    })

    unified_print("开始新闻爬取和AI处理...", "news2docx_cli", "pipeline_processing")

    try:
        # 运行爬虫
        scraper = ScraperNewsScraper(cfg)
        scrape_results = scraper.run()

        # 如果有成功抓取的文章，进行AI处理
        if scrape_results.articles:
            unified_print(f"开始两步AI处理 {len(scrape_results.articles)} 篇新闻...", "news2docx_cli", "pipeline_processing")
            unified_print(f"第一步：词数调整 ({TARGET_WORD_MIN}-{TARGET_WORD_MAX}词)", "news2docx_cli", "pipeline_processing")
            unified_print("第二步：翻译为中文", "news2docx_cli", "pipeline_processing")

            try:
                # 使用新的两步处理方法
                processed_data = process_articles_two_steps_concurrent(scrape_results.articles, target_lang="Chinese")

                article_count = len(processed_data.get('articles', []))
                unified_print(f"两步AI处理完成，处理了 {article_count} 篇新闻", "news2docx_cli", "pipeline_processing")

                # 转换格式以兼容文档生成
                unified_print("正在转换数据格式以兼容文档生成...", "news2docx_cli", "pipeline_processing")
                converted_data = convert_ai_processor_output_to_docx_format(processed_data)
                unified_print(f"格式转换完成，共 {len(converted_data.get('articles', []))} 篇文章", "news2docx_cli", "pipeline_processing")

                # 记录任务结束
                log_task_end("news2docx_cli", "pipeline_processing", True, {
                    "scraped_articles": len(scrape_results.articles),
                    "processed_articles": article_count,
                    "converted_articles": len(converted_data.get('articles', [])),
                    "target_language": "Chinese"
                })

                return converted_data

            except Exception as e:
                handle_error(e, "两步AI处理失败", "news2docx_cli", "pipeline_processing")
                # 返回转换后的原始抓取结果
                unified_print("AI处理失败，返回转换后的原始抓取结果", "news2docx_cli", "pipeline_processing", "warning")

                log_task_end("news2docx_cli", "pipeline_processing", False, {
                    "error": "AI处理失败",
                    "fallback": "转换后的原始抓取结果"
                })

                # 转换Scraper输出格式
                converted_scraper_data = convert_scraper_output_to_docx_format(scrape_results.to_jsonable())
                return converted_scraper_data
        else:
            unified_print("没有成功抓取到任何文章，跳过AI处理", "news2docx_cli", "pipeline_processing", "warning")

            log_task_end("news2docx_cli", "pipeline_processing", False, {
                "error": "无文章可处理"
            })

            # 转换Scraper输出格式
            converted_scraper_data = convert_scraper_output_to_docx_format(scrape_results.to_jsonable())
            return converted_scraper_data

    except NewsProcessingError as e:
        handle_error(e, "新闻处理管道执行失败", "news2docx_cli", "pipeline_processing")
        raise
    except Exception as e:
        handle_error(e, "新闻处理管道执行失败", "news2docx_cli", "pipeline_processing")
        raise NewsProcessingError("新闻处理管道执行失败") from e


# -------------------------------
# DOCX 文档生成部分
# -------------------------------


@dataclass
class FontConfig:
    """字体配置类"""
    name: str
    size_pt: float

    def __post_init__(self) -> None:
        """验证配置参数"""
        if not self.name:
            raise ValueError("字体名称不能为空")
        if self.size_pt <= 0:
            raise ValueError("字体大小必须大于0")


@dataclass
class DocumentConfig:
    """文档配置类"""
    first_line_indent: float
    font_zh: FontConfig
    font_en: FontConfig

    def __post_init__(self) -> None:
        """验证配置参数"""
        if self.first_line_indent < 0:
            raise ValueError("首行缩进不能为负数")


class Language:
    """语言常量"""
    CHINESE = "zh"
    ENGLISH = "en"
    SUPPORTED = (CHINESE, ENGLISH)


class DocumentStyle:
    """文档样式常量"""
    TITLE_SIZE_MULTIPLIER = 1.0
    DEFAULT_TEXT_COLOR = RGBColor(0, 0, 0)
    FILENAME_CLEAN_PATTERN = r'[^\w\s.-]'


class DocumentWriter:
    """
    DOCX文档写入器

    提供专业的文档生成和格式化功能，支持：
    - 中英文双语文档
    - 自定义字体和样式
    - 自动文件名清理
    - 错误处理和验证
    """

    def __init__(self, config: DocumentConfig):
        """
        初始化文档写入器

        Args:
            config: 文档配置对象

        Raises:
            ValueError: 配置参数无效时抛出
        """
        if not isinstance(config, DocumentConfig):
            raise TypeError("config必须是DocumentConfig类型")

        self.config = config
        self.document = Document()
        self._setup_document()

    def _setup_document(self) -> None:
        """设置文档基本属性"""
        # 可以在这里添加更多文档级别的设置
        pass

    def _validate_language(self, language: str) -> None:
        """
        验证语言参数

        Args:
            language: 语言代码

        Raises:
            ValueError: 语言不支持时抛出
        """
        if language not in Language.SUPPORTED:
            supported_str = ", ".join(Language.SUPPORTED)
            raise ValueError(f"不支持的语言 '{language}'，支持的语言: {supported_str}")

    def _get_font_config(self, language: str) -> FontConfig:
        """
        获取指定语言的字体配置

        Args:
            language: 语言代码

        Returns:
            FontConfig: 对应的字体配置
        """
        return self.config.font_zh if language == Language.CHINESE else self.config.font_en

    def _configure_font(self, run, language: str, is_bold: bool = False) -> None:
        """
        配置字体样式

        Args:
            run: docx run对象
            language: 语言代码
            is_bold: 是否加粗
        """
        font_config = self._get_font_config(language)

        run.font.size = Pt(font_config.size_pt)
        run.font.bold = is_bold
        run.font.color.rgb = DocumentStyle.DEFAULT_TEXT_COLOR
        run.font.name = font_config.name

        # 关键部分：同时设置中文和西文字体，确保Word正确识别
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), font_config.name)  # 中文
        rFonts.set(qn("w:ascii"), font_config.name)     # 英文
        rFonts.set(qn("w:hAnsi"), font_config.name)     # 默认西文

    def _clean_filename(self, filename: str) -> str:
        """
        清理文件名，移除不安全的字符

        Args:
            filename: 原始文件名

        Returns:
            str: 清理后的文件名
        """
        if not filename:
            raise ValueError("文件名不能为空")

        # 使用工具函数清理文件名
        cleaned_name = safe_filename(filename)

        # 确保有正确的扩展名
        if not cleaned_name.lower().endswith('.docx'):
            cleaned_name += '.docx'

        return cleaned_name

    def _ensure_directory(self, filepath: Union[str, Path]) -> Path:
        """
        确保目录存在

        Args:
            filepath: 目录路径

        Returns:
            Path: 目录路径对象

        Raises:
            OSError: 创建目录失败时抛出
        """
        try:
            return ensure_directory(filepath)
        except Exception as e:
            raise OSError(f"创建目录失败: {filepath}") from e

    def add_title(self, text: str, language: str) -> None:
        """
        添加标题

        Args:
            text: 标题文本
            language: 语言代码 ('zh' 或 'en')

        Raises:
            ValueError: 参数无效时抛出
        """
        self._validate_language(language)

        if not text:
            raise ValueError("标题文本不能为空")

        # 清理标题文本，移除句号
        cleaned_text = re.sub(r"[。\.]", "", text.strip())

        if not cleaned_text:
            raise ValueError("清理后的标题文本不能为空")

        # 创建标题段落
        title_paragraph = self.document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置标题字体（比正文大一些）
        title_run = title_paragraph.add_run(cleaned_text)
        font_config = self._get_font_config(language)
        title_size = font_config.size_pt * DocumentStyle.TITLE_SIZE_MULTIPLIER

        # 使用统一的字体配置方法，确保中英文字体正确设置
        self._configure_font(title_run, language, is_bold=True)
        title_run.font.size = Pt(title_size)  # 单独覆盖字号（比正文大）

    def add_chinese_title(self, text: str) -> None:
        """
        添加中文标题（宋体居中加粗）

        Args:
            text: 中文标题文本

        Raises:
            ValueError: 文本为空时抛出
        """
        if not text:
            raise ValueError("中文标题文本不能为空")

        # 清理标题文本，移除句号
        cleaned_text = re.sub(r"[。\.]", "", text.strip())

        if not cleaned_text:
            raise ValueError("清理后的中文标题文本不能为空")

        # 创建标题段落
        title_paragraph = self.document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加中文标题文本
        title_run = title_paragraph.add_run(cleaned_text)

        # 计算中文标题字体大小（和正文一样大）
        # 五号字体10.5pt * 标题乘数1.0 = 10.5pt，和正文一样大小
        chinese_title_size = DEFAULT_FONT_SIZE_PT * DEFAULT_TITLE_SIZE_MULTIPLIER

        # 设置中文标题样式：宋体、居中、加粗、大字体
        title_run.font.name = "宋体"  # 中文宋体
        title_run.font.size = Pt(chinese_title_size)  # 使用计算出的字体大小
        title_run.font.bold = True    # 加粗

        # 确保Word正确识别中英文字体
        r = title_run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "宋体")  # 中文
        rFonts.set(qn("w:ascii"), "宋体")     # 英文
        rFonts.set(qn("w:hAnsi"), "宋体")     # 默认西文

    def add_paragraph(self, text: str, language: str) -> None:
        """
        添加段落内容

        Args:
            text: 段落文本（单段内容）
            language: 语言代码 ('zh' 或 'en')

        Raises:
            ValueError: 文本为空或只包含空白字符时抛出
        """
        self._validate_language(language)

        if not text:
            raise ValueError("段落文本不能为空")

        # 清理文本，去除首尾空白字符
        cleaned_text = text.strip()

        if not cleaned_text:
            raise ValueError("清理后的段落文本不能为空")

        # 创建段落
        paragraph = self.document.add_paragraph()

        # 设置首行缩进
        paragraph.paragraph_format.first_line_indent = Inches(self.config.first_line_indent)

        # 添加文本并配置字体
        run = paragraph.add_run(cleaned_text)
        self._configure_font(run, language)

    def add_separator(self) -> None:
        """添加分隔线"""
        separator = self.document.add_paragraph()
        separator.add_run("—" * 50)

    def save(self, filename: str, filepath: Union[str, Path] = ".") -> str:
        """
        保存文档

        Args:
            filename: 文件名
            filepath: 保存路径

        Returns:
            str: 保存的文件完整路径

        Raises:
            ValueError: 文件名无效时抛出
            OSError: 保存失败时抛出
        """
        safe_filename = self._clean_filename(filename)
        directory = self._ensure_directory(filepath)

        full_path = directory / safe_filename

        try:
            self.document.save(str(full_path))
        except Exception as e:
            raise OSError(f"保存文档失败: {full_path}") from e

        return str(full_path)

    def get_document_stats(self) -> dict:
        """
        获取文档统计信息

        Returns:
            dict: 包含段落数、字符数等统计信息的字典
        """
        paragraphs = self.document.paragraphs
        total_chars = sum(len(p.text) for p in paragraphs)

        return {
            "paragraph_count": len(paragraphs),
            "total_characters": total_chars,
            "estimated_pages": max(1, total_chars // 2000)  # 粗略估算
        }


    @staticmethod
    def clean_chinese_filename(filename: str) -> str:
        """
        清理中文文件名，移除不安全的字符并确保文件名有效

        Args:
            filename: 原始文件名（通常是中文标题）

        Returns:
            str: 清理后的文件名
        """
        if not filename:
            return "untitled_document.docx"

        # 使用工具函数清理文件名
        cleaned_name = safe_filename(filename)

        # 确保有正确的扩展名
        if not cleaned_name.lower().endswith('.docx'):
            cleaned_name += '.docx'

        return cleaned_name

    def _process_article_content(self, article: Dict[str, Any]) -> None:
        """
        处理单篇文章内容的私有方法

        Args:
            article: 单篇文章数据
        """
        start_time = time.time()
        # 获取文章ID
        article_id = article.get("id", "unknown")

        input_data = {
            "article_id": article_id,
            "eng_title": article.get("EngTitle", ""),
            "chi_title": article.get("ChiTitle", ""),
            "eng_content_keys": list(article.get("EngContent", {}).keys()),
            "chi_content_keys": list(article.get("ChiContent", {}).keys())
        }

        # 处理英文标题
        eng_title = article.get("EngTitle", "").strip()
        if eng_title:
            log_processing_step("news2docx_cli", "document_creation", f"添加文章 {article_id} 的英文标题", {
                "title": eng_title[:50]
            })
            self.add_title(eng_title, Language.ENGLISH)

        # 处理英文内容段落
        eng_content = article.get("EngContent", {})
        eng_paragraphs_added = 0
        if isinstance(eng_content, dict):
            # 按 p1, p2, p3... 的顺序处理段落
            paragraph_keys = sorted([k for k in eng_content.keys() if k.startswith("p") and k[1:].isdigit()],
                                  key=lambda x: int(x[1:]))
            for key in paragraph_keys:
                content = eng_content[key].strip()
                if content:
                    log_processing_step("news2docx_cli", "document_creation", f"添加文章 {article_id} 的英文段落 {key}")
                    self.add_paragraph(content, Language.ENGLISH)
                    eng_paragraphs_added += 1

        # 处理中文标题
        chi_title = article.get("ChiTitle", "").strip()
        if chi_title:
            log_processing_step("news2docx_cli", "document_creation", f"添加文章 {article_id} 的中文标题", {
                "title": chi_title[:50]
            })
            # 使用专门的中文标题方法（宋体居中加粗）
            self.add_chinese_title(chi_title)

        # 处理中文内容段落
        chi_content = article.get("ChiContent", {})
        chi_paragraphs_added = 0
        if isinstance(chi_content, dict):
            # 按 p1, p2, p3... 的顺序处理段落
            paragraph_keys = sorted([k for k in chi_content.keys() if k.startswith("p") and k[1:].isdigit()],
                                  key=lambda x: int(x[1:]))
            for key in paragraph_keys:
                content = chi_content[key].strip()
                if content:
                    log_processing_step("news2docx_cli", "document_creation", f"添加文章 {article_id} 的中文段落 {key}")
                    self.add_paragraph(content, Language.CHINESE)
                    chi_paragraphs_added += 1

        processing_time = time.time() - start_time
        output_data = {
            "eng_title_added": bool(eng_title),
            "chi_title_added": bool(chi_title),
            "eng_paragraphs_added": eng_paragraphs_added,
            "chi_paragraphs_added": chi_paragraphs_added,
            "processing_time": processing_time
        }

        log_processing_result("news2docx_cli", "document_creation", f"文章 {article_id} 内容处理",
                            input_data, output_data, "success",
                            {"processing_time": processing_time})

    def generate_single_article(self, article: Dict[str, Any]) -> None:
        """
        根据单篇文章数据生成 DOCX 文档内容

        Args:
            article: 单篇文章数据
        """
        try:
            self._process_article_content(article)
        except Exception as e:
            article_id = article.get("id", "unknown") if isinstance(article, dict) else "unknown"
            print(f"处理文章 {article_id} 时发生错误: {e}")

    def generate_from_json(self, json_data: Dict[str, Any]) -> None:
        """
        根据 JSON 数据生成 DOCX 文档

        Args:
            json_data: 包含文章数据的 JSON 字典
        """
        start_time = time.time()
        processed_articles = 0
        failed_articles = 0

        if not json_data or "articles" not in json_data:
            log_error("news2docx_cli", "document_creation", Exception("JSON 数据格式无效"), "缺少 'articles' 字段")
            return

        articles = json_data["articles"]
        if not isinstance(articles, list):
            log_error("news2docx_cli", "document_creation", Exception("'articles' 字段不是列表格式"), "数据格式错误")
            return

        log_processing_step("news2docx_cli", "document_creation", f"开始处理 {len(articles)} 篇文章")

        for i, article in enumerate(articles, 1):
            try:
                # 使用私有方法处理文章内容
                self._process_article_content(article)

                # 在文章之间添加分隔线（除了最后文章）
                if i < len(articles):
                    self.add_separator()

                processed_articles += 1

            except Exception as e:
                article_id = article.get("id", str(i))
                log_error("news2docx_cli", "document_creation", e, f"处理文章 {article_id} 时发生错误")
                failed_articles += 1
                continue

        processing_time = time.time() - start_time

        # 记录批量处理结果
        log_batch_processing(
            "news2docx_cli", "document_creation", "文章内容处理",
            len(articles), processed_articles, failed_articles, processing_time, "completed"
        )

        log_processing_step("news2docx_cli", "document_creation", "JSON 数据处理完成", {
            "total_articles": len(articles),
            "processed_articles": processed_articles,
            "failed_articles": failed_articles,
            "processing_time": processing_time
        })

    @classmethod
    def create_from_json_data(cls, json_data: Dict[str, Any],
                             output_dir: Union[str, Path] = None) -> List[str]:
        """
        从内存中的 JSON 数据为每篇文章创建独立的 DOCX 文档

        Args:
            json_data: JSON 数据字典
            output_dir: 输出目录，如果为None则使用用户桌面的"英语新闻"文件夹

        Returns:
            List[str]: 生成的 DOCX 文件路径列表
        """
        start_time = time.time()

        # 设置默认输出目录为用户桌面的"英语新闻"文件夹
        if output_dir is None:
            import os
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            output_dir = os.path.join(desktop_path, "英语新闻")

        # 确保输出目录存在
        output_path = Path(output_dir)
        try:
            output_path.mkdir(parents=True, exist_ok=True)
            log_processing_step("news2docx_cli", "batch_document_creation", f"输出目录: {output_path}")
        except Exception as e:
            log_error("news2docx_cli", "batch_document_creation", e, "创建输出目录失败")
            return []

        # 验证 JSON 数据
        if not json_data:
            log_error("news2docx_cli", "batch_document_creation", Exception("JSON 数据为空"), "数据验证失败")
            return []

        if "articles" not in json_data or not isinstance(json_data["articles"], list):
            log_error("news2docx_cli", "batch_document_creation", Exception("JSON 数据格式无效"), "缺少有效的 'articles' 字段")
            return []

        articles = json_data["articles"]
        generated_files = []

        log_processing_step("news2docx_cli", "batch_document_creation", f"开始为 {len(articles)} 篇文章生成独立的 DOCX 文档")

        for i, article in enumerate(articles, 1):
            try:
                # 获取中文标题作为文件名
                chi_title = article.get("ChiTitle", "").strip()
                if not chi_title:
                    # 如果没有中文标题，使用英文标题
                    eng_title = article.get("EngTitle", "").strip()
                    if eng_title:
                        chi_title = eng_title
                    else:
                        # 如果都没有标题，使用默认名称
                        article_id = article.get("id", str(i))
                        chi_title = f"文章_{article_id}"

                # 清理文件名
                filename = cls.clean_chinese_filename(chi_title)

                log_processing_step("news2docx_cli", "batch_document_creation", f"正在处理第 {i} 篇文章", {
                    "title": chi_title,
                    "article_id": article.get("id", str(i))
                })

                # 创建新的文档写入器实例
                config = create_default_config()
                writer = cls(config)

                # 生成单篇文章内容
                writer.generate_single_article(article)

                # 保存文档
                file_save_start = time.time()
                try:
                    output_file_path = writer.save(filename, output_path)
                    file_save_time = time.time() - file_save_start

                    # 获取文档统计信息
                    stats = writer.get_document_stats()

                    # 记录文件操作结果
                    log_file_operation(
                        "news2docx_cli", "batch_document_creation", "保存DOCX文档",
                        output_file_path, stats.get("total_characters", 0), file_save_time, "success",
                        {
                            "article_id": article.get("id", str(i)),
                            "paragraph_count": stats.get("paragraph_count", 0),
                            "estimated_pages": stats.get("estimated_pages", 0)
                        }
                    )

                    log_processing_step("news2docx_cli", "batch_document_creation", f"文档已生成", {
                        "file_path": output_file_path,
                        "article_id": article.get("id", str(i)),
                        "file_size": stats.get("total_characters", 0),
                        "save_time": file_save_time
                    })

                    generated_files.append(output_file_path)

                except Exception as e:
                    file_save_time = time.time() - file_save_start

                    log_file_operation(
                        "news2docx_cli", "batch_document_creation", "保存DOCX文档",
                        filename, 0, file_save_time, "error",
                        {"error": str(e), "article_id": article.get("id", str(i))}
                    )

                    log_error("news2docx_cli", "batch_document_creation", e, f"保存文档失败 - 文章 {article.get('id', str(i))}")
                    continue

            except Exception as e:
                article_id = article.get("id", str(i))
                log_error("news2docx_cli", "batch_document_creation", e, f"处理文章 {article_id} 时发生错误")
                continue

        # 记录完整的批量处理结果
        log_batch_processing(
            "news2docx_cli", "batch_document_creation", "批量DOCX文档生成",
            len(articles), len(generated_files), len(articles) - len(generated_files),
            time.time() - start_time, "completed"
        )

        log_processing_step("news2docx_cli", "batch_document_creation", "批量文档生成完成", {
            "total_files": len(generated_files),
            "output_directory": str(output_path)
        })

        return generated_files


def main() -> None:
    """命令行入口函数，解析参数并运行完整流程"""
    import Scraper
    ap = Scraper.build_arg_parser()
    args = ap.parse_args()

    cfg = ScrapeConfig(
        output_dir=args.output_dir,
        api_url=args.api_url,
        api_token=args.api_token,
        max_urls=max(1, args.max_urls),
        concurrency=max(1, args.concurrency),
        retry_interval_hours=max(1, args.retry_hours),
        request_timeout=max(5, args.timeout),
        strict_success=bool(args.strict_success),
        max_api_rounds=max(1, args.max_api_rounds),
        per_url_retries=max(0, args.per_url_retries),
        pick_mode=args.pick_mode,
        random_seed=args.random_seed,
    )

    run_complete_pipeline(cfg)


def create_default_config() -> DocumentConfig:
    """
    创建默认文档配置

    Returns:
        DocumentConfig: 默认配置对象
    """
    return DocumentConfig(
        first_line_indent=0.3,
        font_zh=FontConfig(name="宋体", size_pt=10.5),
        font_en=FontConfig(name="Cambria", size_pt=10.5)
    )


# 为了向后兼容，提供旧的类名
ArticleProcessor = DocumentWriter


# -------------------------------
# 主程序入口
# -------------------------------


def print_header() -> None:
    """打印程序头部信息"""
    print(MAIN_SEPARATOR)
    print("新闻爬虫与文档生成系统")
    print(MAIN_SEPARATOR)
    unified_print("新闻爬虫与文档生成系统启动", "news2docx_cli", "system_init")


def run_document_generation(json_data: Dict[str, Any]) -> List[str]:
    """运行文档生成流程"""
    log_processing_step("news2docx_cli", "document_generation", "开始生成DOCX文档")
    unified_print("开始生成DOCX文档...", "news2docx_cli", "document_generation")

    generated_files = DocumentWriter.create_from_json_data(json_data=json_data)
    return generated_files


def display_results(generated_files: List[str]) -> None:
    """显示处理结果"""
    unified_print("处理结果:", "news2docx_cli", "result_display")

    if generated_files:
        unified_print(f"处理完成！成功生成 {len(generated_files)} 个 DOCX 文档", "news2docx_cli", "result_display")
        unified_print("生成的文件:", "news2docx_cli", "result_display")
        for i, file_path in enumerate(generated_files, 1):
            unified_print(f"{i}. {file_path}", "news2docx_cli", "result_display")
    else:
        unified_print("文档生成失败，请检查错误信息", "news2docx_cli", "result_display", "error")


def handle_system_errors(func):
    """系统级错误处理装饰器"""
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except KeyboardInterrupt:
            print("\n\n⚠️  用户中断执行")
            sys.exit(1)
        except ImportError as e:
            print(f"\n❌ 模块导入错误: {e}")
            print("请确保所有依赖都已正确安装")
            sys.exit(1)
        except NewsProcessingError as e:
            handle_error(e, "新闻处理系统错误")
            sys.exit(1)
        except Exception as e:
            handle_error(e, "程序执行出错")
            sys.exit(1)
    return wrapper


@handle_system_errors
def run_complete_pipeline(cfg: Optional[ScrapeConfig] = None) -> None:
    """运行完整的新闻处理流程：爬取 -> AI处理 -> 文档生成"""
    # 记录主任务开始
    log_task_start("news2docx_cli", "complete_pipeline", {
        "pipeline_steps": ["新闻爬取", "AI处理", "文档生成"]
    })

    print_header()

    # 使用默认配置或提供的配置
    cfg = cfg or ScrapeConfig()

    # 第一步：运行新闻爬虫和AI处理管道，返回处理后的数据
    processed_data = run_news_pipeline(cfg)

    # 第二步：生成DOCX文档（如果有文章被处理）
    if processed_data and processed_data.get("articles"):
        generated_files = run_document_generation(processed_data)
        display_results(generated_files)

        # 记录主任务结束
        log_task_end("news2docx_cli", "complete_pipeline", True, {
            "total_steps": 3,
            "generated_files_count": len(generated_files),
            "processed_articles": len(processed_data.get("articles", []))
        })
    else:
        unified_print("没有成功抓取到文章，跳过文档生成", "news2docx_cli", "complete_pipeline", "warning")

        # 记录主任务结束
        log_task_end("news2docx_cli", "complete_pipeline", False, {
            "error": "无文章数据",
            "completed_steps": 1
        })


# -------------------------------
# Scraper集成功能
# -------------------------------



# -------------------------------
# Scraper集成功能
# -------------------------------

def run_with_scraper(cfg: Optional[ScrapeConfig] = None) -> None:
    """
    使用Scraper模块进行新闻爬取，然后在News2Docx_CLI.py中进行AI处理

    Args:
        scraper_cfg: Scraper模块的配置对象，如果为None则使用默认配置
    """
    print("=" * 60)
    print("使用Scraper模块进行新闻爬取")
    print("=" * 60)

    try:
        # 动态导入Scraper模块
        import Scraper

        # 使用默认配置或提供的配置
        cfg = cfg or Scraper.ScrapeConfig()

        print("🔍 开始使用Scraper模块进行新闻爬取...")
        print("-" * 40)

        # 创建Scraper实例（纯爬虫模式）
        scraper = ScraperNewsScraper(cfg)

        # 运行纯爬虫
        scrape_results = scraper.run()

        print(f"\n✅ 爬取完成！成功抓取 {len(scrape_results.articles)} 篇文章")

        # 保存爬取数据为JSON文件
        if scrape_results.success > 0:
            timestamp = now_stamp()
            save_scraped_data_to_json(scrape_results, timestamp)

        # 将爬取结果传递给News2Docx_CLI.py的AI处理流程
        if scrape_results.articles:
            print(f"\n开始AI处理 {len(scrape_results.articles)} 篇新闻...")

            try:
                # 使用新的两步AI处理流程
                processed_data = process_articles_two_steps_concurrent(scrape_results.articles, target_lang="Chinese")

                print(f"[成功] 两步AI处理完成，处理了 {len(processed_data.get('articles', []))} 篇新闻")

                # 转换格式以兼容文档生成
                print("[信息] 正在转换数据格式以兼容文档生成...")
                converted_data = convert_ai_processor_output_to_docx_format(processed_data)
                print(f"[成功] 格式转换完成，共 {len(converted_data.get('articles', []))} 篇文章")

                # 生成DOCX文档
                if converted_data and converted_data.get("articles"):
                    generated_files = run_document_generation(converted_data)
                    display_results(generated_files)
                else:
                    print("\n[警告] 没有成功处理到文章，跳过文档生成")
                    print(MAIN_SEPARATOR)

            except Exception as e:
                handle_error(e, "AI处理失败")
                # 返回转换后的原始抓取结果
                print("[备选] AI处理失败，返回转换后的原始抓取结果")
                converted_scraper_data = convert_scraper_output_to_docx_format(scrape_results.to_jsonable())
                if converted_scraper_data and converted_scraper_data.get("articles"):
                    generated_files = run_document_generation(converted_scraper_data)
                    display_results(generated_files)
        else:
            print("\n[警告] 没有成功抓取到任何文章，跳过AI处理")
            print(MAIN_SEPARATOR)

    except NewsProcessingError as e:
        handle_error(e, "Scraper模块执行失败")
        raise
    except Exception as e:
        handle_error(e, "Scraper模块执行失败")
        raise NewsProcessingError("Scraper模块执行失败") from e


def build_scraper_arg_parser() -> argparse.ArgumentParser:
    """构建Scraper专用参数解析器"""
    p = argparse.ArgumentParser(description="使用Scraper模块进行新闻爬取（纯爬虫模式），然后进行AI处理")
    p.add_argument("--output-dir", default=os.getenv("CRAWLER_OUTPUT_DIR", ""))
    p.add_argument("--api-url", default=os.getenv("CRAWLER_API_URL", None))
    p.add_argument("--api-token", default=os.getenv("CRAWLER_API_TOKEN", None))
    p.add_argument("--max-urls", type=int, default=int(os.getenv("CRAWLER_MAX_URLS", "10")))
    p.add_argument("--concurrency", type=int, default=int(os.getenv("CRAWLER_CONCURRENCY", str(DEFAULT_CONCURRENCY))))
    p.add_argument("--retry-hours", type=int, default=int(os.getenv("CRAWLER_RETRY_HOURS", "24")))
    p.add_argument("--timeout", type=int, default=int(os.getenv("CRAWLER_TIMEOUT", "30")))
    p.add_argument("--strict-success", action="store_true", default=True)
    p.add_argument("--max-api-rounds", type=int, default=int(os.getenv("CRAWLER_MAX_API_ROUNDS", "5")))
    p.add_argument("--per-url-retries", type=int, default=int(os.getenv("CRAWLER_PER_URL_RETRIES", "2")))
    p.add_argument("--pick-mode", choices=["fifo", "random"], default=os.getenv("CRAWLER_PICK_MODE", "random"),
                   help="URL 取样模式：fifo 或 random（默认 random）")
    p.add_argument("--random-seed", type=int, default=(int(os.getenv("CRAWLER_RANDOM_SEED")) if os.getenv("CRAWLER_RANDOM_SEED") else None),
                   help="随机种子，用于 random 模式下结果可复现")
    p.add_argument("--use-scraper", action="store_true", help="使用Scraper模块进行爬取（AI处理仍在News2Docx_CLI.py中）")

    return p


def main_with_scraper() -> None:
    """使用Scraper模块的命令行入口函数"""
    ap = build_scraper_arg_parser()
    args = ap.parse_args()

    # 动态导入Scraper模块
    import Scraper

    cfg = Scraper.ScrapeConfig(
        output_dir=args.output_dir,
        api_url=args.api_url,
        api_token=args.api_token,
        max_urls=max(1, args.max_urls),
        concurrency=max(1, args.concurrency),
        retry_interval_hours=max(1, args.retry_hours),
        request_timeout=max(5, args.timeout),
        strict_success=bool(args.strict_success),
        max_api_rounds=max(1, args.max_api_rounds),
        per_url_retries=max(0, args.per_url_retries),
        pick_mode=args.pick_mode,
        random_seed=args.random_seed,
    )

    run_with_scraper(cfg)


def test_font_sizes():
    """测试字体大小设置"""
    print("测试字体大小设置...")

    try:
        # 计算各种字体大小
        base_size = DEFAULT_FONT_SIZE_PT  # 五号字体 = 10.5pt
        title_multiplier = DEFAULT_TITLE_SIZE_MULTIPLIER  # 1.0
        chinese_title_size = base_size * title_multiplier  # 10.5 * 1.0 = 10.5pt

        print(f"✅ 正文字体大小（五号）: {base_size}pt")
        print(f"✅ 标题放大乘数: {title_multiplier}")
        print(f"✅ 英文标题大小: {base_size * title_multiplier:.1f}pt (和正文一样大)")
        print(f"✅ 中文标题大小: {chinese_title_size:.1f}pt (和正文一样大)")

        # 创建配置并测试
        config = create_default_config()
        writer = DocumentWriter(config)

        # 测试英文标题
        writer.add_title("English Title Test", "en")

        # 测试中文标题
        writer.add_chinese_title("中文标题测试")

        print("✅ 字体大小设置测试完成")
        return True

    except Exception as e:
        print(f"❌ 字体大小测试失败: {e}")
        return False


def test_format_conversion():
    """测试格式转换功能"""
    print("测试格式转换功能...")

    # 测试ai_processor格式转换
    ai_processor_sample = {
        "metadata": {
            "total_articles": 1,
            "successful_articles": 1,
            "processing_method": "two_step_processing"
        },
        "articles": [
            {
                "id": "1",
                "original_title": "Test English Title",
                "translated_title": "测试英文标题",
                "original_content": "This is the first paragraph.\nThis is the second paragraph.",
                "adjusted_content": "This is the expanded first paragraph with more content.\nThis is the expanded second paragraph with additional details.",
                "adjusted_word_count": 450,
                "translated_content": "这是第一个段落，经过扩展的内容。\n这是第二个段落，包含更多细节。",
                "target_language": "Chinese",
                "processing_timestamp": "20241201_120000",
                "url": "https://example.com/test",
                "success": True
            }
        ]
    }

    # 测试Scraper格式转换
    scraper_sample = {
        "total": 1,
        "success": 1,
        "failed": 0,
        "articles": [
            {
                "index": 1,
                "url": "https://example.com/test",
                "title": "Test Title",
                "content": "This is test content.\nSecond paragraph here.",
                "content_length": 100,
                "word_count": 20,
                "scraped_at": "20241201_120000"
            }
        ],
        "successful_urls": ["https://example.com/test"],
        "failed_urls": []
    }

    try:
        # 测试ai_processor转换
        converted_ai = convert_ai_processor_output_to_docx_format(ai_processor_sample)
        print(f"✅ AI处理器格式转换成功: {len(converted_ai.get('articles', []))} 篇文章")

        # 测试Scraper转换
        converted_scraper = convert_scraper_output_to_docx_format(scraper_sample)
        print(f"✅ Scraper格式转换成功: {len(converted_scraper.get('articles', []))} 篇文章")

        # 验证转换结果格式
        if converted_ai.get('articles') and len(converted_ai['articles']) > 0:
            article = converted_ai['articles'][0]
            required_keys = ['id', 'EngTitle', 'ChiTitle', 'EngContent', 'ChiContent']
            missing_keys = [key for key in required_keys if key not in article]
            if missing_keys:
                print(f"❌ AI转换结果缺少必需字段: {missing_keys}")
            else:
                print("✅ AI转换结果格式正确")

        print("格式转换测试完成!")
        return True

    except Exception as e:
        print(f"❌ 格式转换测试失败: {e}")
        return False


if __name__ == "__main__":
    # 检查是否有命令行参数
    if len(sys.argv) > 1:
        # 检查是否使用Scraper模块进行爬取
        if "--use-scraper" in sys.argv:
            # 使用Scraper模块进行爬取，AI处理仍在News2Docx_CLI.py中
            main_with_scraper()
        elif "--test-conversion" in sys.argv:
            # 测试格式转换功能
            test_format_conversion()
        elif "--test-fonts" in sys.argv:
            # 测试字体大小设置
            test_font_sizes()
        else:
            # 使用内置爬取功能
            main()
    else:
        # 没有命令行参数，直接运行完整流程
        run_complete_pipeline()